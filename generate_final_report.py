from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_final_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('Wireless Weightlifting Trolley - Final Project Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. Project Overview
    doc.add_heading('1. Project Overview', level=1)
    doc.add_paragraph(
        "This project involves the design and development of a wireless, remotely controlled weightlifting machine (Trolley). "
        "The system uses high-torque 24V DC worm gear motors for lifting and a 24V 15A SMPS for stable power delivery. "
        "The core control is handled by an ESP32 microcontroller, allowing for mobile application integration."
    )

    # 2. Final Component List
    doc.add_heading('2. System Components', level=1)
    
    components = [
        "Primary Power: 24V 15A 360W SMPS (Regulated Switching Power Supply).",
        "Control Brain: ESP32 Development Board (WiFi + Bluetooth).",
        "Lifting Motors: 24V 27RPM High Torque DC Worm Gear Motors (Self-locking).",
        "Motor Drivers: 24V High-Current H-Bridges (Recommended: BTS7960 43A).",
        "Power Management: DC-DC Buck Converter (24V to 5V for ESP32/Relays).",
        "Safety/Logic: 4-Channel Relay Module (High-power safety interlocks).",
        "Wiring: 18 Gauge Power Wires (15m) and Logic Jumper Wires."
    ]
    for comp in components:
        doc.add_paragraph(comp, style='List Bullet')

    # 3. Hardware Wiring Map
    doc.add_heading('3. Hardware Wiring & Connections', level=1)
    
    connections = [
        "Power Input: SMPS 24V (+) and (-) to Motor Driver Power Inputs.",
        "Logic Power: SMPS 24V (+) to Buck Converter Input (+). Buck Converter Output set to 5V.",
        "ESP32 Power: Buck Converter 5V Output to ESP32 'VIN' pin and 'GND'.",
        "Control Signals: ESP32 GPIO 12 to Motor Driver R_PWM (Forward/Up).",
        "Control Signals: ESP32 GPIO 13 to Motor Driver L_PWM (Backward/Down).",
        "Common Ground: Ensure all GND pins (ESP32, Buck Converter, Motor Driver) are connected together."
    ]
    for conn in connections:
        doc.add_paragraph(conn, style='List Bullet')

    # 4. Programming Setup (Arduino IDE)
    doc.add_heading('4. Software Setup Instructions', level=1)
    setup_steps = [
        "Install Arduino IDE from arduino.cc.",
        "Go to File > Preferences. Add ESP32 URL: https://raw.githubusercontent.com/espressif/arduino-esp32/gh-pages/package_esp32_index.json",
        "Go to Tools > Board > Boards Manager. Search for 'ESP32' and Install.",
        "Select Board: 'DOIT ESP32 DEVKIT V1'.",
        "Ensure the Micro-USB cable supports data transfer (not just charging)."
    ]
    for step in setup_steps:
        doc.add_paragraph(step, style='List Number')

    # 5. ESP32 Testing Code
    doc.add_heading('5. Prototype Testing Code (Web Control)', level=1)
    code_desc = doc.add_paragraph("The following code creates a WiFi Hotspot named 'Trolley_Control'. Access the control panel via 192.168.4.1 in any mobile browser.")
    
    # Add code block with a distinct font
    code_text = (
        "#include <WiFi.h>\n"
        "#include <WebServer.h>\n\n"
        "const int motorForward = 12; // UP\n"
        "const int motorBackward = 13; // DOWN\n"
        "WebServer server(80);\n\n"
        "void setup() {\n"
        "  pinMode(motorForward, OUTPUT);\n"
        "  pinMode(motorBackward, OUTPUT);\n"
        "  WiFi.softAP(\"Trolley_Control\", \"12345678\");\n"
        "  server.on(\"/\", [](){ server.send(200, \"text/html\", \"<h1>Trolley Active</h1><a href='/up'>UP</a><br><a href='/down'>DOWN</a><br><a href='/stop'>STOP</a>\"); });\n"
        "  server.on(\"/up\", [](){ digitalWrite(motorForward, HIGH); digitalWrite(motorBackward, LOW); server.send(200); });\n"
        "  server.on(\"/down\", [](){ digitalWrite(motorForward, LOW); digitalWrite(motorBackward, HIGH); server.send(200); });\n"
        "  server.on(\"/stop\", [](){ digitalWrite(motorForward, LOW); digitalWrite(motorBackward, LOW); server.send(200); });\n"
        "  server.begin();\n"
        "}\n\n"
        "void loop() { server.handleClient(); }"
    )
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    # 6. Future Improvements (Safety & Feedback)
    doc.add_heading('6. Safety & Future-Proofing Recommendations', level=1)
    futures = [
        "Limit Switches: To stop the lift automatically at the top and bottom.",
        "Ultrasonic Sensors: For autonomous obstacle avoidance.",
        "Load Cell (HX711): To display the weight being lifted on the mobile app.",
        "Current Sensing: To detect motor overload and shut down automatically."
    ]
    for future in futures:
        doc.add_paragraph(future, style='List Bullet')

    # Save
    file_name = 'Weightlifting_Trolley_Final_Report.docx'
    doc.save(file_name)
    print(f"Final report saved as {file_name}")

if __name__ == "__main__":
    create_final_report()
