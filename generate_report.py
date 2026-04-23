from docx import Document
from docx.shared import Inches

def create_project_report():
    doc = Document()
    doc.add_heading('Wireless Weightlifting Machine (Trolley) - Project Report', 0)

    # 1. Project Overview
    doc.add_heading('1. Project Overview', level=1)
    doc.add_paragraph(
        "The project aims to develop a wireless weightlifting trolley remotely controlled by a mobile application. "
        "The machine is designed to lift and transport heavy weights, providing a versatile and efficient solution for mechanical material handling."
    )

    # 2. Existing Assets (Found in USEFULL Folder)
    doc.add_heading('2. Existing Project Assets', level=1)
    doc.add_paragraph("Based on the technical assets provided:")
    
    assets = [
        "Mechanical Design: CAD model showing a 4-wheel chassis with an integrated lifting mechanism (pulley and belt driven).",
        "Motors: 24V 27RPM High Torque DC Worm Gear Motors (Model 5840-31ZY). Self-locking feature ensures safety during lifting.",
        "Power Source: 24V 15A 360W SMPS Regulated Switching Power Supply.",
        "Current Prototype State: Hardware assembly and motor setup are completed."
    ]
    for asset in assets:
        doc.add_paragraph(asset, style='List Bullet')

    # 3. Required Components (To be Added)
    doc.add_heading('3. Required Components for Integration', level=1)
    doc.add_paragraph("The following components are essential to complete the wireless control and power management system:")
    
    missing_components = [
        "4x Relays: For high-power switching or as safety interlocks.",
        "Buck Converter: To step down 24V from SMPS to 5V/3.3V for the control logic.",
        "ESP32 Microcontroller: For WiFi/Bluetooth connectivity and mobile app interface.",
        "24V Motor Drivers (e.g., BTS7960 43A): Capable of handling the high stall current (3.1A) of the motors.",
        "Wiring: 18 Gauge Wires (15m) for power and Jumper wires for logic signals.",
        "15A SMPS: (Identified, confirmed for procurement)."
    ]
    for comp in missing_components:
        doc.add_paragraph(comp, style='List Bullet')

    # 4. Technical Suggestions for College Prototype
    doc.add_heading('4. Strategic Suggestions for Success', level=1)
    
    suggestions = [
        "Motor Selection & Control: Use High-Current H-Bridges (BTS7960) to avoid overheating during heavy lifts. Ensure the motor driver can handle the stall torque requirements.",
        "Safety Interlocks: Install Limit Switches at the top and bottom of the lifting mechanism to prevent mechanical failure from over-travel.",
        "Power Isolation: Use the Buck converter with proper filtering to prevent electrical noise from the motors from resetting the ESP32.",
        "Remote Interface: Develop the mobile application using Blynk IoT or a local Web Server hosted on the ESP32 for low-latency control.",
        "Feedback Systems: Consider adding a Load Cell (HX711) to display the weight being lifted on the mobile app—this will be highly impressive for faculty presentations.",
        "Emergency Stop: Implement both a physical E-Stop and a software 'failsafe' that stops all motors if the WiFi connection is lost."
    ]
    for suggestion in suggestions:
        doc.add_paragraph(suggestion, style='List Bullet')

    # 5. Presentation Advice
    doc.add_heading('5. Presentation Tips for Faculty', level=1)
    doc.add_paragraph(
        "Focus on the 'Autonomous/Remote' aspect. Highlight the self-locking nature of worm gear motors as a key safety feature. "
        "Demonstrate the power efficiency of using a regulated 24V system over standard 12V."
    )

    # Save the document
    file_name = 'Wireless_Weightlifting_Trolley_Report.docx'
    doc.save(file_name)
    print(f"Report saved as {file_name}")

if __name__ == "__main__":
    create_project_report()
